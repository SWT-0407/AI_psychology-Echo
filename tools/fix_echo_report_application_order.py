from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement


TARGET = Path(r"E:\Users\醨\PycharmProjects\PythonProject6\reports\心语Echo项目报告_情感关系应用强化版.docx")


def find_paragraph(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text[:80]}")


def insert_before(paragraph, text: str, style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if style:
        new_para.style = style
    new_para.text = text
    return new_para


def remove_between(doc: Document, start_text: str, end_text: str) -> None:
    start_i = end_i = None
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == start_text:
            start_i = i
        elif text == end_text and start_i is not None:
            end_i = i
            break
    if start_i is None or end_i is None:
        raise ValueError("Could not locate range to remove")
    for paragraph in list(doc.paragraphs[start_i + 1:end_i]):
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def main() -> None:
    doc = Document(TARGET)
    remove_between(doc, "七、项目应用前景和社会价值", "八、项目存在的问题及今后的改进方向")
    before_eight = find_paragraph(doc, "八、项目存在的问题及今后的改进方向")
    section_items = [
        ("（一）从心理工具扩展为情感关系基础设施", "Heading 2"),
        ("Echo 的应用前景不应只停留在校园心理支持。它更大的价值，是成为现代人处理情感表达、关系连接、记忆保存和自我理解的 AI 情感基础设施。用户真正需要的往往不是一个会说话的聊天框，而是一个能承接“我有话没说完”“我不知道该不该联系 TA”“我想保存一段关系”“我想和过去的自己和解”的空间。", "Normal"),
        ("在这个定位下，Echo 的用户不再只是“有心理困扰的人”，而是所有在现代生活中出现情感封闭、关系困惑、表达困难、怀念与告别需求的人。心理评测、AI 树洞、虚拟角色和用户画像只是入口，真正的长期价值是让用户能把碎片化情绪转化为可理解、可回看、可行动的关系与成长线索。", "Normal"),
        ("（二）未完成对话与关系决策支持", "Heading 2"),
        ("很多现实痛苦来自未完成对话：没和父母解释清楚，没和朋友道歉，没向喜欢的人表达，没和分开的人告别，没把委屈和边界说出口。Echo 可以把这类需求从单纯倾诉升级为关系决策支持。用户输入想说但没说的话后，系统先识别对话类型，例如告白、道歉、解释误会、表达委屈、提出边界、请求帮助、感谢、告别或复合。", "Normal"),
        ("随后，系统评估这段对话是否值得完成、是否适合发送、现实沟通风险有多高，以及用户此刻的情绪是否适合立刻行动。Echo 可以给出“发送、暂缓、改写、不发送、仪式化告别、寻求第三方支持”等建议，并生成温和表达版、清晰边界版、简短不纠缠版、道歉修复版和给自己看的整理版。这样，系统不是鼓励用户逃避真实关系，而是帮助用户以更安全、更清楚的方式回到现实沟通。", "Normal"),
        ("（三）记忆回声与情感告别场景", "Heading 2"),
        ("在强流动社会中，很多关系会突然断裂：亲人离世、朋友疏远、毕业分别、分手、离开家乡、搬离熟悉的城市。Echo 可以发展“记忆回声”模式，帮助用户把聊天记录、照片、语音片段、共同事件和自己的回忆整理成一份关系记忆档案。系统可以基于这些材料生成“记忆式回应”，让用户把未说出口的话安放下来。", "Normal"),
        ("这个场景的边界必须非常清楚：Echo 不是复活某个人，也不能冒充真实个体。它做的是记忆整理、告别支持和情感缓冲。对于失去重要关系的人来说，价值不只是“再聊一次”，而是帮助用户完成遗憾整理、关系纪念、情绪封存和后续生活重建。", "Normal"),
        ("（四）关系预演、crush 与授权人格陪伴", "Heading 2"),
        ("Echo 还可以进入关系预演与理想关系互动场景。用户在表白、拒绝、道歉、求助、面试、家庭沟通前，可以先和系统模拟一次。系统结合用户画像判断用户是害怕冲突、害怕被讨厌、过度自责，还是缺少表达结构，再给出更合适的话术与下一步应对。", "Normal"),
        ("对于 crush、idol、虚拟角色和二次元陪伴需求，Echo 可以提供两类合规路径：一类是授权 IP 或授权人格陪伴，用于粉丝经济和品牌互动；另一类是灵感型理想关系角色，只模拟某种气质、说话风格和关系感，不冒充具体真人。这个方向的商业潜力很强，但必须坚持身份标识、授权边界和防沉迷机制，让虚拟关系帮助用户理解真实需求，而不是替代现实关系。", "Normal"),
        ("（五）人生转场、自我对话与成长档案", "Heading 2"),
        ("人生转场是 Echo 很适合切入的高频场景：大一入学、毕业、考研失败、就业焦虑、刚入职、分手、搬家、离开家乡、成为照护者或进入新的家庭角色。人在这些阶段不是单纯压力大，而是旧身份正在结束，新身份还没有稳定。Echo 可以提供阶段性转场陪伴：回顾、告别、整理、重建和行动计划。", "Normal"),
        ("基于长期用户画像，Echo 可以让用户和过去的自己、未来的自己对话：三个月前的自己为什么痛苦，现在的自己发生了什么变化，一年后的自己会如何安慰现在的自己。系统也可以生成情绪时间胶囊、关系地图、低能量生活计划和个人成长档案，让用户看见自己不是孤立地崩溃，而是在具体关系、事件和阶段中变化。", "Normal"),
        ("（六）校园、家庭与公共服务价值", "Heading 2"),
        ("在校园场景中，Echo 仍然可以作为心理咨询中心的前置整理工具。学生端保留树洞、日记、陪伴、未完成对话和个人报告；咨询中心端只在明确授权或匿名聚合条件下查看趋势，不直接窥探原始聊天。它能帮助学生在正式咨询前把近期睡眠、压力、人际关系、家庭沟通和情绪变化说清楚，降低首次求助门槛。", "Normal"),
        ("在家庭和亲密关系场景中，Echo 可以把用户复杂的情绪翻译成不同对象能理解的表达版本：给父母看的版本、给朋友看的版本、给咨询师看的版本、给自己的整理版。它的社会价值不在于用 AI 替代真实关系，而在于修复真实表达的断点，降低求助羞耻感，帮助用户把“我说不清楚哪里不对”转化成更容易沟通的状态线索。", "Normal"),
        ("（七）未来商业化与行业前景", "Heading 2"),
        ("Echo 的未来商业化不一定依赖大规模收集隐私数据，而可以围绕本地部署、校园授权版、隐私合规评估、专业内容库、授权 IP 陪伴、人生档案服务和情感记忆保存展开。它可以从大学生心理陪伴切入，逐步扩展到青年关系预演、粉丝/IP 情感互动、家庭沟通辅助、人生转场陪伴、老年人生档案和心理咨询前画像系统。", "Normal"),
        ("最终，Echo 可以形成的不是一个单点应用，而是一套面向大众情感需求的 AI 关系智能体平台：处理和自己的关系、和现实他人的关系、和想象对象的关系、和过去记忆的关系。它的关键不是制造虚拟依赖，而是在合适的边界内，为现代人的情感封闭、内求需求和关系断裂提供新的承接方式。", "Normal"),
    ]
    for text, style in section_items:
        insert_before(before_eight, text, style)
    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
