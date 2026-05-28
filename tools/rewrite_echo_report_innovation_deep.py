from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement


SRC = Path(r"E:\Users\醨\PycharmProjects\PythonProject6\reports\心语Echo项目报告_情感关系应用强化版.docx")
OUT = Path(r"E:\Users\醨\PycharmProjects\PythonProject6\reports\心语Echo项目报告_技术创新强化版.docx")


def find_paragraph(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text[:80]}")


def insert_before(paragraph, text: str, style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if style:
        new_para.style = style
    new_para.text = text
    return new_para


def remove_between(doc: Document, start_text: str, end_text: str) -> None:
    start_i = end_i = None
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == start_text:
            start_i = i
        elif text == end_text and start_i is not None:
            end_i = i
            break
    if start_i is None or end_i is None:
        raise ValueError("Could not locate range to remove")
    for paragraph in list(doc.paragraphs[start_i + 1:end_i]):
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(SRC)

    remove_between(doc, "六、项目创新点", "七、项目应用前景和社会价值")
    before_seven = find_paragraph(doc, "七、项目应用前景和社会价值")

    items = [
        ("（一）构建 EchoMind-DIPS 动态对话智能处理模型", "Heading 2"),
        ("本项目不是简单调用大模型接口，而是构建了 EchoMind-DIPS Model（Dialogue Intelligence Processing System）作为心理对话的总控模型。该模型把用户自然输入拆解为“多模态理解—画像检索—动态 Prompt 构建—大模型推理—安全干预—结构化输出—评分更新—记忆回写”八个环节，使系统能够在一次对话中同时完成理解、追问、陪伴、评估和长期画像更新。", "Normal"),
        ("EchoMind-DIPS 的创新点在于把普通聊天流程改造成心理场景中的闭环推理流程。系统会根据用户当前六维状态、历史画像、最近情绪、关系对象和安全风险动态生成对话策略，而不是每轮都把用户当成第一次出现的人。这样既能保留大模型自然表达能力，又能让回复受到心理评估、长期记忆和安全边界的共同约束。", "Normal"),
        ("（二）构建 SEEN 六维心理状态评分模型", "Heading 2"),
        ("项目构建了 SEEN Model（Scientific Echo Evaluation Network）作为六维心理状态评分模型。模型以 X=[x1,x2,x3,x4,x5,x6] 表示情绪状态、焦虑控制力、生理状态、行为与动力、社交与支持、认知与意义，并通过特征归一化、线性心理主干、维度交互层、深层心理模式层、风险调节层和稳定化输出层生成 0-100 的状态评分。", "Normal"),
        ("SEEN 的价值不只是“打分”，而是把自然语言里的复杂心理线索转化为可解释、可追踪、可进入画像系统的结构化状态。传统问卷通常只得到一次性结果，普通情绪分类只判断开心或难过；SEEN 则关注情绪、压力、睡眠、行动力、支持系统和意义感之间的联动关系，并通过风险调节层避免高风险表达被平均分掩盖。", "Normal"),
        ("（三）构建 EchoMemory 长时记忆与时序情感图谱", "Heading 2"),
        ("项目构建 EchoMemory 作为长时记忆与用户画像中枢。它不只是保存聊天记录，而是把日记、树洞、评测、角色陪伴、多模态输入和安全判断统一整理为带来源、时间戳、置信度、证据片段和衰减权重的画像信号。系统进一步将人物、事件、时间、情绪强度、压力源、缓解方式和关系对象组织为时序情感知识图谱。", "Normal"),
        ("这一设计使 Echo 具备长期状态理解能力。系统可以区分一次性情绪波动和反复出现的结构性压力源，也可以发现用户在某类关系、某个时间段或某种场景下反复出现的情绪模式。后续还可接入 Temporal Graph Network、TGAT、DyGFormer 或动态异构图 Transformer，用图神经网络对长期情感轨迹进行趋势预测和关系推理。", "Normal"),
        ("（四）提出用户画像驱动的功能一体化架构", "Heading 2"),
        ("Echo 的创新不在于功能数量多，而在于所有功能都围绕同一份动态用户画像运转。心理评测负责识别状态，树洞负责承接真实表达，虚拟角色负责延续关系记忆，未完成对话模块负责判断现实沟通策略，RAG 负责补充可靠知识，安全模块负责控制风险边界，最终都回写到画像中枢。", "Normal"),
        ("这种架构让系统从“多个功能页面”升级为“画像驱动的情感智能体”。用户下一次进入系统时，Echo 能根据画像决定是继续追问、温和陪伴、建议现实求助、生成行动计划，还是进入关系预演和记忆整理。画像既是数据层的统一表示，也是模型层的个性化上下文，还是产品层的体验中枢。", "Normal"),
        ("（五）构建 UDR 未完成对话解决模型", "Heading 2"),
        ("针对现代人普遍存在的“有话说不出口”问题，项目设计了 UDR Model（Unfinished Dialogue Resolver）。该模型将未完成对话从单纯倾诉升级为关系决策支持：先识别对话类型，如告白、道歉、解释误会、表达委屈、提出边界、请求帮助、感谢、告别或复合，再判断这段话是否值得完成、是否适合发送、对方是否可能接住，以及现实沟通风险有多高。", "Normal"),
        ("UDR 的技术流程包括意图识别、关系对象建模、完成价值评估、风险成本评估、表达策略生成和可能回应模拟。系统可以输出发送型完成、不发送型完成、仪式型完成、模拟型完成和转介型完成五类建议，并生成温和表达版、清晰边界版、简短不纠缠版、道歉修复版和给自己看的整理版。这个模型让 Echo 不止接住情绪，还能帮助用户更成熟地回到真实关系。", "Normal"),
        ("（六）提出情感人格边界与记忆回声模型", "Heading 2"),
        ("面向怀念的人、crush、idol、虚拟角色和人生档案等场景，项目提出 APBE（Affective Persona Boundary Engine，情感人格边界引擎）与 EchoMemorial 记忆回声模型。系统可以根据用户提供的聊天记录、照片、事件描述、语气偏好和关系回忆生成“记忆式回应”，帮助用户完成遗憾整理、关系纪念和情感告别。", "Normal"),
        ("这一创新的关键不是复刻真人，而是建立清晰边界：对离开或逝去的人，系统标记为记忆生成而非复活；对 idol 或公众人物，系统区分授权人格和灵感型角色；对 crush 场景，系统强调关系预演和自我理解而非冒充对方。技术上，系统通过身份边界标签、授权状态标识、人格相似度约束和防沉迷提示，避免虚拟关系越界。", "Normal"),
        ("（七）构建跨模态情感时序对齐机制", "Heading 2"),
        ("项目在多模态层面提出 CMTA（Cross-modal Temporal Alignment）跨模态情感时序对齐机制。文字、语音、图片和表情不是孤立输入，而是被投影到同一条情绪时间轴中：文本提供语义线索，语音停顿和语速提供压力线索，图片提供情境线索，表情识别提供低权重辅助线索。", "Normal"),
        ("为了避免“摄像头一看就诊断”的误导，CMTA 对多模态输出采用置信度门控、滑动窗口平滑和冲突降权机制。未来可接入 Qwen-VL、Qwen-Audio、Whisper / wav2vec2、CLIP 类视觉语义编码器和跨模态注意力融合模型，使多模态线索只用于辅助理解和趋势观察，而不直接作为临床判断。", "Normal"),
        ("（八）构建 Hybrid RAG / GraphRAG 心理知识增强框架", "Heading 2"),
        ("项目将普通向量检索升级为 Hybrid RAG / GraphRAG 心理知识增强框架。系统不仅使用 bge-small-zh-v1.5 和 Chroma 进行语义召回，还设计了关键词召回、来源可信度评分、适用人群标签、知识版本管理、证据冲突消解和 reranker 重排序流程。", "Normal"),
        ("这一框架的创新点在于，知识库不是简单拼进 prompt，而是先判断资料是否可靠、是否适合大学生或当前关系场景、是否与用户画像冲突。未来可接入 BGE-M3、ColBERT、BGE reranker、GraphRAG 和知识图谱路径检索，让系统在生成心理报告、未完成对话建议和行动计划时具有更强证据基础。", "Normal"),
        ("（九）提出安全校准与伦理约束模型", "Heading 2"),
        ("心理与情感关系场景中，安全边界本身就是核心技术创新。项目设计了 Safety Calibration Gate（安全校准门控），在普通对话前先进行自伤、伤人、极端绝望、操控性关系、沉迷虚拟关系和身份冒充风险识别。系统不仅输出风险等级，还保留触发依据、置信度和建议处置等级。", "Normal"),
        ("技术上，安全模块可进一步接入不确定性校准、conformal prediction、自伤风险专用分类器、安全奖励模型和 RLAIF 安全偏好优化。当模型不确定时，系统不会强行生成有感染力的回答，而是降低生成自由度，转向澄清追问、现实求助提示或更保守的陪伴策略。对记忆回声和 idol 陪伴等场景，系统还加入身份标识、授权校验和虚拟关系边界提示。", "Normal"),
        ("（十）形成本地个性化与隐私计算闭环", "Heading 2"),
        ("Echo 的另一个技术创新，是把个性化能力和用户数据主权放在同一套架构里考虑。系统支持不登录本地运行、授权后云同步、完整聊天内容上传开关、数据导出和删除入口；同时采用本地 LoRA / Adapter 的轻量个性化路线，让用户端保留交互样本和个性化小权重，云端只同步摘要、评分或脱敏后的适配器差分。", "Normal"),
        ("未来该闭环可以继续接入联邦式 LoRA / Adapter 聚合、差分隐私噪声注入、端侧加密索引、敏感字段自动脱敏和可信执行环境。这样，Echo 不需要集中收集大量私密对话，也能逐步学习用户偏好的陪伴方式、关系表达习惯和情绪调节模式。对于心理和情感产品来说，这种“隐私优先的个性化学习”比单纯堆模型参数更有长期价值。", "Normal"),
        ("（十一）建立模型评测与持续迭代机制", "Heading 2"),
        ("项目不是只展示页面效果，而是建立了面向比赛和后续产品化的模型评测闭环。系统可以对比基础大模型、本地 LoRA 模型、RAG 增强模型、SEEN 评分结果和安全门控结果，记录回复质量、风险召回率、RAG 命中率、响应延迟、用户反馈分和画像更新效果。", "Normal"),
        ("这一机制让 Echo 的技术路线具备可验证性。每一次用户反馈、星级评分、画像确认、误判修正和安全触发，都可以成为后续模型优化数据。未来可进一步引入 DPO、RLAIF、主动学习和专家标注复核，让系统在共情、追问、建议、转介、关系预演和记忆回应之间形成更稳定的策略选择能力。", "Normal"),
    ]

    for text, style in items:
        insert_before(before_seven, text, style)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
