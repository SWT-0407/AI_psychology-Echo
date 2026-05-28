from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.text import WD_BREAK


SRC = Path(r"C:\Users\醨\Desktop\附件：人工智能创新赛报名表、查新报告模版.docx")
OUT = Path(r"E:\Users\醨\PycharmProjects\PythonProject6\reports\附件：人工智能创新赛报名表、查新报告_Echo模型补充版.docx")


PROJECT_NAME = "心语 Echo——基于用户画像的 AI 心理陪伴与状态感知系统"
AUTHORS = "孙文韬  徐婧雯"
FINISH_DATE = "2026年5月28日"


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.text = text


def set_cell_text(cell, text: str) -> None:
    cell.text = ""
    parts = text.split("\n")
    if not parts:
        return
    first = cell.paragraphs[0]
    first.text = parts[0]
    for part in parts[1:]:
        p = cell.add_paragraph()
        p.text = part


def fill_cover_fields(doc: Document) -> None:
    replacements = {
        "项目名称：": f"项目名称：{PROJECT_NAME}",
        "项目作者：": f"项目作者：{AUTHORS}",
        "查新完成日期：": f"查新完成日期：{FINISH_DATE}",
    }
    seen = set()
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in replacements and text not in seen:
            paragraph.text = replacements[text]
            seen.add(text)


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(SRC)
    fill_cover_fields(doc)

    # The sixth table is the novelty-search report body in the supplied template.
    table = doc.tables[5]

    set_cell_text(table.rows[0].cells[1], PROJECT_NAME)
    set_cell_text(
        table.rows[1].cells[1],
        "报名参加第28届中国机器人及人工智能大赛人工智能创新赛。通过查新，重点查证本项目在“动态心理评估、用户画像驱动的一体化心理陪伴、六维心理状态评分、长时记忆与个性化反馈、安全边界控制”等方面，是否存在与本项目整体技术路线相同或高度相似的公开研究或产品方案。",
    )
    set_cell_text(
        table.rows[2].cells[1],
        "本项目围绕大学生日常心理记录与陪伴场景，构建了以用户画像为中枢的 AI 心理陪伴与状态感知系统。相较于传统问卷、普通聊天机器人或单一情感陪伴产品，Echo 将日记评测、AI 树洞、虚拟角色、长时记忆、知识增强和安全干预组织成闭环，形成可持续更新的心理状态画像。\n"
        "1. EchoMind-DIPS Model（Dialogue Intelligence Processing System）：面向心理评测与树洞对话的动态对话智能处理模型。系统从用户自然表达出发，经 Qwen 多模态理解、心理支持知识库 RAG、EchoMemory 历史记忆调用、动态 Prompt 构建、DeepSeek 语言推理、安全干预模块、标准化 JSON 输出、SEEN 评分引擎和 EchoMemory 回写，形成“输入—理解—追问—评分—回复—记忆更新”的闭环流程。\n"
        "2. SEEN Model（Scientifecho Evaluation Network）：六维心理状态评分模型。模型以 X=[x1,x2,x3,x4,x5,x6] 表示情绪状态、焦虑控制力、生理状态、行为与动力、社交与支持、认知与意义，经过特征归一化、线性心理主干、维度交互层、深层心理模式层、风险调节层、融合层与稳定化输出，形成 0-100 的六维稳定化评分与综合评估结果。\n"
        "3. EchoMemory Long-term Memory & User Profiling Module：长时记忆与用户画像模块。系统统一存储对话记录、阶段评分、情绪标签、关键事件、用户偏好与长期心理轨迹，支持本地优先、云同步可选、记忆检索与调用、角色记忆管理、历史记录可视化和动态 Prompt 反馈。\n"
        "4. 用户画像驱动的一体化功能闭环：评测、树洞、虚拟陪伴、RAG、风险预警和成长反馈不再是分散功能，而是共同读写同一份动态用户画像，使系统能够根据长期状态调整追问方向、陪伴语气、报告解释和风险响应策略。",
    )
    set_cell_text(
        table.rows[3].cells[1],
        "查新点1：是否已有系统同时面向大学生日常心理记录、AI 树洞、虚拟角色陪伴和用户画像，形成评测—倾诉—陪伴—反馈的一体化闭环。\n"
        "查新点2：是否已有公开方案提出与 EchoMind-DIPS 相同的心理对话处理流程，即多模态理解、RAG 知识增强、长时记忆调用、动态 Prompt 构建、安全干预、标准化输出、评分引擎和记忆回写的闭环式心理评测模型。\n"
        "查新点3：是否已有公开模型采用与 SEEN Model 相同的六维心理状态输入、特征归一化、线性主干、维度交互、深层心理模式、风险调节和稳定化输出结构，用于动态评估大学生心理状态变化。\n"
        "查新点4：是否已有公开系统将 EchoMemory 类长时记忆模块作为用户画像中枢，统一管理对话记录、阶段评分、情绪标签、关键事件、用户偏好、角色记忆和长期心理轨迹，并反向驱动个性化陪伴与报告生成。\n"
        "查新点5：是否已有公开产品在心理陪伴场景中同时强调本地优先存储、云同步可选、完整聊天内容上传开关、高风险表达安全门控和用户可控的隐私边界。",
    )
    set_cell_text(
        table.rows[4].cells[1],
        "文献检索范围：Google Scholar、CNKI/知网、万方、IEEE Xplore、ACM Digital Library、PubMed、JMIR、arXiv、Semantic Scholar，以及 Woebot、Wysa 等公开产品资料和搜索引擎公开页面。\n"
        "检索时间范围：2017年1月至2026年5月。\n"
        "中文检索词：心理健康聊天机器人、AI 心理陪伴、大学生心理评估、动态心理评估、心理用户画像、长期记忆、情绪树洞、虚拟陪伴、六维心理模型、危机干预、心理 RAG。\n"
        "英文检索词：mental health chatbot, AI mental health companion, conversational agent mental health, digital mental health intervention, long-term memory chatbot, user profiling mental health, personalized mental health support, LLM mental health safety, psychological assessment model, RAG mental health, crisis detection chatbot.\n"
        "组合检索式示例：(mental health chatbot OR AI mental health companion) AND (long-term memory OR user profile OR longitudinal tracking); (psychological assessment OR mental health screening) AND (conversational agent OR chatbot) AND (six-dimensional OR multidimensional); (LLM mental health) AND (safety OR crisis detection OR ethical safeguard).",
    )
    set_cell_text(
        table.rows[5].cells[1],
        "按上述检索词在公开数据库和网络资料中检索，发现已有相关研究和产品主要集中在以下方向：\n"
        "1. Woebot 等心理健康聊天机器人研究表明，基于 CBT 框架的自动对话代理可用于年轻人的抑郁、焦虑自助干预，但其核心更偏向结构化 CBT 内容推送与短周期对话干预，并未形成面向大学生日常记录、虚拟角色陪伴、用户画像和长时记忆的一体化系统。\n"
        "2. Wysa 等产品和相关研究关注 AI 聊天、心理教育、放松训练和部分人工教练协同，能够提供情绪支持和自助练习，但公开资料未见其采用与 EchoMind-DIPS 相同的“多模态理解—RAG—动态 Prompt—安全干预—SEEN 评分—EchoMemory 回写”的闭环模型。\n"
        "3. 近年关于大语言模型在心理健康中的综述指出，LLM 可用于心理健康问题识别、支持性对话和危机信号检测，但也普遍存在安全边界、可解释性、长期跟踪和评估体系不足等问题。Echo 针对这些问题引入风险门控、标准化输出、用户画像和长期轨迹分析。\n"
        "4. 现有心理量表和心理评估系统大多采用固定题项或单次问卷，虽然具有较强标准化优势，但不适合持续捕捉用户自然表达中的长期变化。SEEN Model 将自然对话中的状态线索转换为六维心理状态评分，并通过风险调节层控制输出稳定性，区别于普通量表打分和单纯情绪分类。\n"
        "5. 现有情感陪伴产品通常强调角色互动和拟人化体验，但较少把角色记忆、心理评分、风险识别、知识增强和用户可控隐私统一到同一画像中枢。EchoMemory 的特点在于把长期心理轨迹、角色互动记忆和个性化建议共同纳入可检索、可回看、可反馈的记忆结构。\n"
        "综上，检索到的相关研究和产品与本项目在部分功能上存在交叉，例如心理聊天、情绪记录、危机提示或 AI 陪伴等，但未见与本项目三个自建模型 EchoMind-DIPS、SEEN Model、EchoMemory 及其一体化闭环完全相同的公开报道。",
    )
    set_cell_text(
        table.rows[6].cells[1],
        "经对检索出的相关文献和公开产品资料进行分析、对比，结论如下：\n"
        "文献与产品1类：以 Woebot、Wysa 等为代表的心理健康聊天机器人，主要提供 CBT、心理教育、冥想训练、情绪支持或人工教练协同，证明了自动对话代理在心理健康支持中的可行性，但公开资料未见其具备 Echo 所设计的六维心理状态评分模型、长时用户画像中枢和虚拟角色陪伴闭环。\n"
        "文献与产品2类：大语言模型心理健康应用研究关注共情回复、问题识别、风险检测和伦理安全，但多停留在单轮或短程对话生成、检测任务或评测框架层面，未见与 EchoMind-DIPS 相同的多模块工程流程和标准化状态输出结构。\n"
        "文献与产品3类：心理评估和情绪识别研究通常针对问卷量表、文本情绪分类或风险识别，缺少将日记、树洞、虚拟陪伴、多模态输入和长期记忆统一为动态用户画像的产品化方案。\n"
        "本项目的主要新颖性体现在：①提出 EchoMind-DIPS 动态对话智能处理模型，将多模态理解、RAG、历史记忆、动态 Prompt、安全干预、标准化输出和评分引擎组织成闭环；②提出 SEEN 六维心理状态评分模型，将自然表达映射为可解释的六维稳定化评分；③提出 EchoMemory 长时记忆与用户画像模块，实现本地优先、云同步可选、长期心理轨迹分析和角色记忆管理；④以用户画像为中枢，把心理评测、AI 树洞、虚拟陪伴、风险预警和成长反馈整合为一体化心理支持系统。\n"
        "检索中未见与本项目整体技术路线和三类自建模型组合完全相同的公开研究或产品报道。因此，本项目在面向大学生日常心理支持的多模态状态感知、用户画像驱动的一体化陪伴、动态心理评估和长时记忆反馈方面具有一定新颖性和应用推广价值。",
    )
    set_cell_text(
        table.rows[8].cells[1],
        "1. 项目 PPT《Echo——基于用户画像的 AI 心理陪伴与状态感知系统》；\n"
        "2. 项目研究报告《心语 Echo 项目报告》；\n"
        "3. 项目源代码与模型说明材料；\n"
        "4. 相关公开文献与产品检索记录。"
    )
    set_cell_text(
        table.rows[9].cells[1],
        "本查新报告根据项目 PPT、项目报告、源代码结构和公开文献检索结果整理。Echo 不是医疗诊断系统，不替代心理咨询师；项目所称心理状态评分、用户画像和风险识别均用于日常自我记录、陪伴支持和求助引导。"
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
