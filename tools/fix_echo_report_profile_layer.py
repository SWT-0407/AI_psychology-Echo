from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement


TARGET = Path(r"E:\Users\醨\PycharmProjects\PythonProject6\reports\心语Echo项目报告_画像价值前景强化版.docx")


def find_paragraph(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text[:80]}")


def insert_after(paragraph, text: str, style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if style:
        new_para.style = style
    new_para.text = text
    return new_para


def main() -> None:
    doc = Document(TARGET)
    anchor = find_paragraph(
        doc,
        "状态建模部分把对话评分转化为六维健康分，再计算综合指数、功能受损等级和风险保护门控。报告展示给用户的是可理解的自然语言；后台结构化结果则进入画像中枢和时序心理状态图谱，用于历史趋势、画像更新、主动陪伴计划、RAG 检索约束和后续追问。",
    )
    h = insert_after(anchor, "（四）画像中枢层：多源信号融合与功能调度", "Heading 2")
    p = insert_after(
        h,
        "画像中枢层负责把来自评测、树洞、虚拟角色、日记、多模态输入和历史记录的信号统一成结构化画像。每条画像信号都会记录来源、时间、置信度、相关事件和可解释证据，既能被用户回看，也能被模型作为个性化上下文使用。",
        "Normal",
    )
    p = insert_after(
        p,
        "在功能调度上，画像中枢会决定系统下一步更适合做什么：如果用户只是短暂低落，系统优先给出陪伴和轻量行动建议；如果连续多天出现睡眠、动力和社交下降，系统会提高追问密度并生成趋势报告；如果出现危机线索，系统会把普通聊天切换为安全回复和现实求助引导。这样功能之间不是平行堆叠，而是由同一个画像状态统一调度。",
        "Normal",
    )
    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
