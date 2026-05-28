from pathlib import Path

from docx import Document


TARGET = Path(r"E:\Users\醨\PycharmProjects\PythonProject6\reports\附件：人工智能创新赛报名表、查新报告_Echo模型补充版.docx")


def set_cell_text(cell, text: str) -> None:
    cell.text = ""
    lines = text.split("\n")
    cell.paragraphs[0].text = lines[0] if lines else ""
    for line in lines[1:]:
        cell.add_paragraph(line)


def main() -> None:
    doc = Document(TARGET)
    table = doc.tables[5]

    row_text = {
        1: (
            "一．查新目的\n"
            "报名参加第28届中国机器人及人工智能大赛人工智能创新赛。通过查新，重点查证本项目在“动态心理评估、用户画像驱动的一体化心理陪伴、六维心理状态评分、长时记忆与个性化反馈、安全边界控制”等方面，是否存在与本项目整体技术路线相同或高度相似的公开研究或产品方案。"
        ),
        2: (
            "二．查新项目的创新要点\n"
            "本项目围绕大学生日常心理记录与陪伴场景，构建了以用户画像为中枢的 AI 心理陪伴与状态感知系统。相较于传统问卷、普通聊天机器人或单一情感陪伴产品，Echo 将日记评测、AI 树洞、虚拟角色、长时记忆、知识增强和安全干预组织成闭环，形成可持续更新的心理状态画像。\n"
            "1. EchoMind-DIPS Model（Dialogue Intelligence Processing System）：面向心理评测与树洞对话的动态对话智能处理模型。系统从用户自然表达出发，经 Qwen 多模态理解、心理支持知识库 RAG、EchoMemory 历史记忆调用、动态 Prompt 构建、DeepSeek 语言推理、安全干预模块、标准化 JSON 输出、SEEN 评分引擎和 EchoMemory 回写，形成“输入—理解—追问—评分—回复—记忆更新”的闭环流程。\n"
            "2. SEEN Model（Scientifecho Evaluation Network）：六维心理状态评分模型。模型以 X=[x1,x2,x3,x4,x5,x6] 表示情绪状态、焦虑控制力、生理状态、行为与动力、社交与支持、认知与意义，经过特征归一化、线性心理主干、维度交互层、深层心理模式层、风险调节层、融合层与稳定化输出，形成 0-100 的六维稳定化评分与综合评估结果。\n"
            "3. EchoMemory Long-term Memory & User Profiling Module：长时记忆与用户画像模块。系统统一存储对话记录、阶段评分、情绪标签、关键事件、用户偏好与长期心理轨迹，支持本地优先、云同步可选、记忆检索与调用、角色记忆管理、历史记录可视化和动态 Prompt 反馈。\n"
            "4. 用户画像驱动的一体化功能闭环：评测、树洞、虚拟陪伴、RAG、风险预警和成长反馈不再是分散功能，而是共同读写同一份动态用户画像，使系统能够根据长期状态调整追问方向、陪伴语气、报告解释和风险响应策略。"
        ),
        3: (
            "三．查新点\n"
            "查新点1：是否已有系统同时面向大学生日常心理记录、AI 树洞、虚拟角色陪伴和用户画像，形成评测—倾诉—陪伴—反馈的一体化闭环。\n"
            "查新点2：是否已有公开方案提出与 EchoMind-DIPS 相同的心理对话处理流程，即多模态理解、RAG 知识增强、长时记忆调用、动态 Prompt 构建、安全干预、标准化输出、评分引擎和记忆回写的闭环式心理评测模型。\n"
            "查新点3：是否已有公开模型采用与 SEEN Model 相同的六维心理状态输入、特征归一化、线性主干、维度交互、深层心理模式、风险调节和稳定化输出结构，用于动态评估大学生心理状态变化。\n"
            "查新点4：是否已有公开系统将 EchoMemory 类长时记忆模块作为用户画像中枢，统一管理对话记录、阶段评分、情绪标签、关键事件、用户偏好、角色记忆和长期心理轨迹，并反向驱动个性化陪伴与报告生成。\n"
            "查新点5：是否已有公开产品在心理陪伴场景中同时强调本地优先存储、云同步可选、完整聊天内容上传开关、高风险表达安全门控和用户可控的隐私边界。"
        ),
        4: (
            "四．文献检索范围及检索策略\n"
            "文献检索范围：Google Scholar、CNKI/知网、万方、IEEE Xplore、ACM Digital Library、PubMed、JMIR、arXiv、Semantic Scholar，以及 Woebot、Wysa 等公开产品资料和搜索引擎公开页面。\n"
            "检索时间范围：2017年1月至2026年5月。\n"
            "中文检索词：心理健康聊天机器人、AI 心理陪伴、大学生心理评估、动态心理评估、心理用户画像、长期记忆、情绪树洞、虚拟陪伴、六维心理模型、危机干预、心理 RAG。\n"
            "英文检索词：mental health chatbot, AI mental health companion, conversational agent mental health, digital mental health intervention, long-term memory chatbot, user profiling mental health, personalized mental health support, LLM mental health safety, psychological assessment model, RAG mental health, crisis detection chatbot。\n"
            "组合检索式示例：(mental health chatbot OR AI mental health companion) AND (long-term memory OR user profile OR longitudinal tracking); (psychological assessment OR mental health screening) AND (conversational agent OR chatbot) AND (six-dimensional OR multidimensional); (LLM mental health) AND (safety OR crisis detection OR ethical safeguard)."
        ),
        5: (
            "五．检索结果\n"
            "按上述检索词在公开数据库和网络资料中检索，发现已有相关研究和产品主要集中在以下方向：\n"
            "1. Fitzpatrick 等在 JMIR Mental Health 发表的 Woebot 随机对照试验显示，基于 CBT 的自动对话代理可用于大学生抑郁、焦虑自助干预；该系统以短周期 CBT 对话、心情跟踪和活动建议为主，未见其形成面向日记评测、虚拟角色、长时用户画像和多模型评分的一体化架构。\n"
            "2. Inkster 等关于 Wysa 的真实世界数据研究显示，AI 情绪支持聊天机器人可用于数字心理健康和用户参与分析；Wysa 公开资料强调情绪支持、自助练习、隐私和临床安全控制，但未见其采用与 EchoMind-DIPS、SEEN、EchoMemory 相同的组合模型。\n"
            "3. 2023 年 JMIR 对心理健康会话代理干预随机对照试验的系统综述和 meta 分析表明，聊天机器人、虚拟人、VR 会话代理等对部分心理健康指标有一定作用，但研究对象多为特定干预任务或症状改善评估，较少涉及跨功能用户画像中枢和长期心理轨迹建模。\n"
            "4. 2023 年 npj Digital Medicine 关于 AI 会话代理促进心理健康与幸福感的系统综述和 meta 分析表明，AI 会话代理在心理健康照护中具有潜力，但现有研究通常围绕干预有效性、用户评价和心理结果指标展开，未见与本项目整体闭环完全一致的模型流程。\n"
            "5. 2024 年以来关于大语言模型在心理健康中的系统综述指出，LLM 可用于心理健康问题识别、支持性对话和危机信号检测，但仍存在安全边界、长期跟踪、可解释性和临床责任等挑战。Echo 针对这些不足引入风险门控、标准化输出、用户画像和长期轨迹分析。\n"
            "综上，检索到的相关研究和产品与本项目在部分功能上存在交叉，例如心理聊天、情绪记录、危机提示或 AI 陪伴等，但未见与本项目三个自建模型 EchoMind-DIPS、SEEN Model、EchoMemory 及其一体化闭环完全相同的公开报道。"
        ),
        6: (
            "六．查新结论\n"
            "经对检索出的相关文献和公开产品资料进行分析、对比，结论如下：\n"
            "文献与产品1类：以 Woebot、Wysa 等为代表的心理健康聊天机器人，主要提供 CBT、心理教育、冥想训练、情绪支持或人工教练协同，证明了自动对话代理在心理健康支持中的可行性，但公开资料未见其具备 Echo 所设计的六维心理状态评分模型、长时用户画像中枢和虚拟角色陪伴闭环。\n"
            "文献与产品2类：大语言模型心理健康应用研究关注共情回复、问题识别、风险检测和伦理安全，但多停留在单轮或短程对话生成、检测任务或评测框架层面，未见与 EchoMind-DIPS 相同的多模块工程流程和标准化状态输出结构。\n"
            "文献与产品3类：心理评估和情绪识别研究通常针对问卷量表、文本情绪分类或风险识别，缺少将日记、树洞、虚拟陪伴、多模态输入和长期记忆统一为动态用户画像的产品化方案。\n"
            "本项目的主要新颖性体现在：①提出 EchoMind-DIPS 动态对话智能处理模型，将多模态理解、RAG、历史记忆、动态 Prompt、安全干预、标准化输出和评分引擎组织成闭环；②提出 SEEN 六维心理状态评分模型，将自然表达映射为可解释的六维稳定化评分；③提出 EchoMemory 长时记忆与用户画像模块，实现本地优先、云同步可选、长期心理轨迹分析和角色记忆管理；④以用户画像为中枢，把心理评测、AI 树洞、虚拟陪伴、风险预警和成长反馈整合为一体化心理支持系统。\n"
            "检索中未见与本项目整体技术路线和三类自建模型组合完全相同的公开研究或产品报道。因此，本项目在面向大学生日常心理支持的多模态状态感知、用户画像驱动的一体化陪伴、动态心理评估和长时记忆反馈方面具有一定新颖性和应用推广价值。"
        ),
        7: (
            "七．申报者本人、所在学院签字盖章的查新声明与证明\n"
            "报告中陈述的事实是真实和准确的。\n"
            "我们按照大赛查新规范进行查新、文献分析和审核，并做出上述查新结论。\n"
            "申报者（签字）：                      申报者所在学院（盖章）：\n"
            "年     月     日"
        ),
        8: (
            "八．附件清单\n"
            "1. 项目 PPT《Echo——基于用户画像的 AI 心理陪伴与状态感知系统》；\n"
            "2. 项目研究报告《心语 Echo 项目报告》；\n"
            "3. 项目源代码与模型说明材料；\n"
            "4. 相关公开文献与产品检索记录。"
        ),
        9: (
            "九．备注\n"
            "本查新报告根据项目 PPT、项目报告、源代码结构和公开文献检索结果整理。Echo 不是医疗诊断系统，不替代心理咨询师；项目所称心理状态评分、用户画像和风险识别均用于日常自我记录、陪伴支持和求助引导。"
        ),
    }

    for row_index, text in row_text.items():
        set_cell_text(table.rows[row_index].cells[0], text)

    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
